#!/usr/bin/env python3
"""EndoCode: build .docx from markdown-like lines. Requires: pip install python-docx"""
from __future__ import annotations

import json
import sys
from pathlib import Path


def main() -> int:
    if len(sys.argv) < 2:
        print("usage: gen_docx.py config.json", file=sys.stderr)
        return 1
    cfg = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    out = Path(cfg["out"])
    title = (cfg.get("title") or out.stem or "Document").strip()
    md = cfg.get("markdown") or cfg.get("content") or ""

    try:
        from docx import Document
    except ImportError:
        print("Brak biblioteki: pip install python-docx", file=sys.stderr)
        return 2

    doc = Document()
    doc.add_heading(title, 0)
    for line in md.splitlines():
        raw = line.rstrip()
        if not raw.strip():
            continue
        s = raw.strip()
        if s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=2)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=1)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=0)
        elif s.startswith(("- ", "* ")):
            doc.add_paragraph(s[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(s)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(str(out))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
